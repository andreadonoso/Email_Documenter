import os.path
import datetime
import base64
import docx 
import re
from io import StringIO
from bs4 import BeautifulSoup
from dotenv import load_dotenv

from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError


SCOPES = ["https://www.googleapis.com/auth/gmail.readonly", "https://www.googleapis.com/auth/calendar.readonly", 'https://www.googleapis.com/auth/calendar.events']


# The file token.json stores the user's access and refresh tokens, and is created
# automatically when the authorization flow completes for the first time.
def authorize(creds):
  if os.path.exists("token.json"):
    creds = Credentials.from_authorized_user_file("token.json", SCOPES)
    
  # If there are no (valid) credentials available, let the user log in.
  if not creds or not creds.valid:
    if creds and creds.expired and creds.refresh_token:
      creds.refresh(Request())
    else:
      flow = InstalledAppFlow.from_client_secrets_file(
          "credentials.json", SCOPES
      )
      creds = flow.run_local_server(port=0)
      
    # Save the credentials for the next run
    with open("token.json", "w") as token:
      token.write(creds.to_json())
      
  return creds

# Removes unwanted extra spaces with a single space or maximum 2 new lines
def normalizeText(text):
  text = re.sub(r'[͏\u200c\xa0\u200d]+', '', text)
  text = re.sub(r'[ \t\r\f\v]*\n[ \t\r\f\v]*', '\n', text)
  text = re.sub(r'\n{3,}', '\n\n', text)
  return text

# Decodes the raw message to visible text for multi-part (from plain, html, x-amp-html mimeTypes) messages and single-part messages
def decodeBody(fullMsg):
  # Checks for multi-part messages
  if 'parts' in fullMsg['payload']:
    numParts = len(fullMsg['payload']['parts'])
    # print(f"NUM PARTS: {numParts}")
    
    visibleTextParts = []
    for count, part in enumerate(fullMsg['payload']['parts'], start=1):
      mimeType = part['mimeType']
      bodyData = part['body'].get('data', None)
      
      if not bodyData:
        continue 
  
      content = base64.urlsafe_b64decode(bodyData).decode('utf-8')
      soup = BeautifulSoup(content, 'html.parser')
      text = normalizeText(soup.get_text(separator="\n", strip=True))
      # return text
      visibleTextParts.append(text) 
    visibleText = "\n\n* * * * * * * * * * * * * * * * * * * * * * * * *  PART CHANGE  * * * * * * * * * * * * * * * * * * * * * * * * *\n\n".join(visibleTextParts)
    return  visibleText
  # Checks for single-part messages       
  else:
    mimeType = fullMsg['payload']['mimeType']
    
    if mimeType == 'text/plain': 
      visibleText = base64.urlsafe_b64decode(fullMsg['payload']['body']['data']).decode('utf-8')
      return normalizeText(visibleText)
    elif mimeType == 'text/html':
      html_content = base64.urlsafe_b64decode(fullMsg['payload']['body']['data']).decode('utf-8')
      soup = BeautifulSoup(html_content, 'html.parser')
      visibleText = soup.get_text(separator="\n", strip=True)
      return normalizeText(visibleText)

# Gets each email's decoded body content and update the queryRes['messages'] to have the fully decoded body ('fullMessage')
# Example: queryRes['messages'][0]['fullMessage'] contains the decoded body of the first message in the message list
def decodeResults(queryRes, service_gmail):
  if 'messages' not in queryRes or (queryRes['resultSizeEstimate'] <= 0):
    return 
  messages = queryRes['messages']
  
  for message in messages:
    fullMessage = service_gmail.users().messages().get(userId="me", id=message['id'], format="full").execute()
    fullMessage['payload']['visibleText'] = decodeBody(fullMessage)
    message['fullMessage'] = fullMessage

# Performs the email search query with the Gmail API
def performSearchQuery(service_gmail, sentFrom, label, subject, body, numResults):
  f = f'from:"{sentFrom}"' if sentFrom else ''
  l = f'label:"{label}"' if label else ''
  s = f'subject:"{subject}"' if subject else ''
  searchQuery = f'{f} {l} {s} {body}' if any([f, l, s, body]) else 'All mail'
  
  # Call the Gmail API 
  return searchQuery, service_gmail.users().messages().list(userId="me", q=searchQuery, maxResults=numResults).execute()

# Decodes the messages and displays them
def showQueryResults(searchQuery, queryRes, numResults):
  print("\n\n-----------------------------------------------------------------------------------------------------------------------------------------------------------------\n\n")
  print(f"Query:\t {searchQuery.strip()}\n")
  print(f"Result size estimate:\t{queryRes['resultSizeEstimate']}")
  print(f"Max results to show:\t{numResults}\n")
  
  if 'messages' not in queryRes or (queryRes['resultSizeEstimate'] <= 0):
    print("No emails found\n")
    print("\n-----------------------------------------------------------------------------------------------------------------------------------------------------------------\n\n")
    return 
  messages = queryRes['messages']
  
  for count,message in enumerate(messages, start=1):
    print(f"\n\n\n{count}) - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - - -\n\n\n")
    
    headers = message['fullMessage']['payload']['headers']
    resSubject = next((header['value'] for header in headers if header['name'] == 'Subject'), "No Subject")
    print(f"SUBJECT:  {resSubject}\n")
    
    visibleText = message['fullMessage']['payload']['visibleText']
    if visibleText:
        print("\nBODY:\n")
        print(f"{visibleText}")
    else:
        print("No visible text found.")
  print("\n-----------------------------------------------------------------------------------------------------------------------------------------------------------------\n\n")

# Writes decoded messages to a word document
def writeToWordDoc(queryRes):
  if 'messages' not in queryRes or (queryRes['resultSizeEstimate'] <= 0):
    return 
  messages = queryRes['messages']
  
  load_dotenv()
  directory = os.getenv("MY_DIRECTORY")
  if not os.path.isdir(directory):
      print(f"Invalid directory: {directory}")
      exit()
      
  today = datetime.datetime.now()
  file_name = "3rdPartyMaintenance_" + today.strftime("%m-%d-%y") + ".docx" 
  path = os.path.join(directory, file_name)
  doc = docx.Document() 
  
  print(f"WRITING TO DOC. PLEASE WAIT . . .")
  for count,message in enumerate(messages, start=1):
    headers = message['fullMessage']['payload']['headers']
    resSubject = next((header['value'] for header in headers if header['name'] == 'Subject'), "No Subject")
    
    doc.add_heading(resSubject + "\n", level=2).bold = True
    
    visibleText = message['fullMessage']['payload']['visibleText']
    if visibleText:
        doc.add_paragraph(visibleText)
    else:
        doc.add_paragraph("No visible text found.")
        
    doc.add_page_break()
    
  print(f"\n\nSAVING DOCUMENT TO PATH:  {path}\n\n")
  doc.save(path) 
  print("Document saved.\n")

# Creates a calendar event based on the body of an email
def createEvents(service_gcal, queryRes):
  if 'messages' not in queryRes or (queryRes['resultSizeEstimate'] <= 0):
    return 
  messages = queryRes['messages']
  
  numEvents = 0
  if(queryRes['resultSizeEstimate'] > 0):
    print(f"\nCREATING EVENTS. PLEASE WAIT . . .\n\n")
  
    for count,message in enumerate(messages, start=1):
      visibleText = message['fullMessage']['payload']['visibleText']

      patterns = {
        "summary": r'Summary:\s*"(.*?)"',
        "description": r'Description:\s*"(.*?)"',
        "startDateTime": r'Start date time:\s*"(.*?)"',
        "startTimeZone": r'Start time zone:\s*"(.*?)"',
        "endDateTime": r'End date time:\s*"(.*?)"',
        "endTimeZone": r'End time zone:\s*"(.*?)"',
      }

      parts = visibleText.split("* * * * * * * * * * * * * * * * * * * * * * * * *  PART CHANGE  * * * * * * * * * * * * * * * * * * * * * * * * *")
      part1 = parts[0]
      
      data = {}
      for key, pattern in patterns.items():
          match = re.search(pattern, part1)
          if match:
              data[key] = match.group(1)
          
      summary = data.get('summary')
      description = data.get('description')
      startDateTime = data.get('startDateTime')
      startTimeZone = data.get('startTimeZone')
      endDateTime = data.get('endDateTime')
      endTimeZone = data.get('endTimeZone')
      
      event = {
        'summary': summary,
        'location': '',
        'description': description,
        'start': {
          'dateTime': startDateTime,
          'timeZone': startTimeZone,
        },
        'end': {
          'dateTime': endDateTime,
          'timeZone': endTimeZone,
        },
        'recurrence': [
        ],
        'attendees': [
        ],
        'reminders': {
          'useDefault': True,
          'overrides': [
          ],
        },
      }

      event = service_gcal.events().insert(calendarId='primary', body=event).execute()
      print('Event created: %s' % (event.get('htmlLink')))
      
      if(count == len(messages)):
        numEvents = count
    
  print(f"\nEvents created: {numEvents}.\n\n")

def main():
  """
    Uses the Gmail API to search for emails and then prints them to a word document. 
    It creates Google Calendar events using the Google Calendar API based on the email event information.
  """
  # AUTHORIZATION
  creds = None
  creds = authorize(creds)

  try:
    sentFrom = "donosoandrea@icloud.com"
    label = ""
    subject = "Test Email with Event"
    body = ""
    numResults = 10  # Max Results = 500
  
    # PERFORM SEARCH QUERY
    service_gmail = build("gmail", "v1", credentials=creds) # Gmail API service object 
    searchQuery, queryRes = performSearchQuery(service_gmail, sentFrom, label, subject, body, numResults)
    decodeResults(queryRes, service_gmail)
    
    # SHOW SEARCH QUERY RESULTS
    showQueryResults(searchQuery, queryRes, numResults)
    
    # TBD: FILTER VALID EMAILS
    
    # WRITE VALID EMAILS TO WORD DOC
    writeToWordDoc(queryRes)
    
    # CREATE CALENDAR EVENTS
    service_gcal = build("calendar", "v3", credentials=creds)
    createEvents(service_gcal, queryRes)

  except HttpError as error:
    # TODO(developer) - Handle errors from gmail API.
    print(f"An error occurred: {error}")


if __name__ == "__main__":
  main()        
